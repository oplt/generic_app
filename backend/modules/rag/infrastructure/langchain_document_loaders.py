from __future__ import annotations

import asyncio
import logging
import tempfile
from pathlib import Path

from backend.modules.rag.domain.models import ParsedDocument

logger = logging.getLogger(__name__)


def _parse_txt_or_md(content: bytes) -> list[ParsedDocument]:
    text = content.decode("utf-8", errors="replace").strip()
    if not text:
        return []
    return [ParsedDocument(content=text, metadata={"format": "text"})]


def _parse_csv(content: bytes) -> list[ParsedDocument]:
    import csv
    import io

    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = [" | ".join(cell.strip() for cell in row if cell.strip()) for row in reader]
    joined = "\n".join(row for row in rows if row)
    if not joined:
        return []
    return [ParsedDocument(content=joined, metadata={"format": "csv"})]


def _parse_pdf(content: bytes) -> list[ParsedDocument]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required for PDF parsing. Install with: uv add pypdf") from exc

    import io

    reader = PdfReader(io.BytesIO(content))
    docs: list[ParsedDocument] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            docs.append(
                ParsedDocument(
                    content=text,
                    metadata={"format": "pdf", "page_number": index},
                    page_number=index,
                )
            )
    return docs


def _parse_docx(content: bytes) -> list[ParsedDocument]:
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX parsing. Install with: uv add python-docx"
        ) from exc

    import io

    document = docx.Document(io.BytesIO(content))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    if not text:
        return []
    return [ParsedDocument(content=text, metadata={"format": "docx"})]


def _parse_with_langchain_loader(file_path: str, loader_name: str) -> list[ParsedDocument]:
    """Fallback LangChain community loaders when available."""
    try:
        if loader_name == "text":
            from langchain_community.document_loaders import TextLoader

            loader = TextLoader(file_path, encoding="utf-8")
        elif loader_name == "csv":
            from langchain_community.document_loaders import CSVLoader

            loader = CSVLoader(file_path)
        else:
            return []
        lc_docs = loader.load()
        return [
            ParsedDocument(
                content=doc.page_content,
                metadata=dict(doc.metadata),
                page_number=doc.metadata.get("page"),
            )
            for doc in lc_docs
            if doc.page_content.strip()
        ]
    except Exception:
        logger.debug("LangChain loader %s unavailable, using native parser", loader_name)
        return []


def _parse_bytes_sync(
    *,
    content: bytes,
    filename: str,
    content_type: str,
) -> list[ParsedDocument]:
    """Parse raw bytes into documents. CPU-bound — use ``load_documents_from_bytes``."""
    suffix = Path(filename).suffix.lower().lstrip(".")
    if suffix in {"txt", "md", "markdown"} or content_type.startswith("text/"):
        return _parse_txt_or_md(content)
    if suffix == "csv":
        parsed = _parse_csv(content)
        if parsed:
            return parsed
    if suffix == "pdf" or content_type == "application/pdf":
        return _parse_pdf(content)
    if suffix == "docx" or content_type in {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }:
        return _parse_docx(content)

    with tempfile.NamedTemporaryFile(suffix=f".{suffix}", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    try:
        return _parse_with_langchain_loader(tmp_path, suffix)
    finally:
        Path(tmp_path).unlink(missing_ok=True)


async def load_documents_from_bytes(
    *,
    content: bytes,
    filename: str,
    content_type: str,
) -> list[ParsedDocument]:
    """Parse file bytes into ParsedDocument list off the event loop thread."""
    return await asyncio.to_thread(
        _parse_bytes_sync,
        content=content,
        filename=filename,
        content_type=content_type,
    )
